from django.shortcuts import render, redirect
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views import View
from django.http import HttpResponse
from sal9event import settings
from .models import Acara as AcaraModel, Bagian as BagianModel, Tugas as TugasModel
from django.core.mail import send_mail
import docx
import calendar;
import time
# Create your views here.


class Acara(LoginRequiredMixin, View):
    login_url = settings.LOGOUT_REDIRECT_URL
    template_name = 'acara.html'

    def get(self, request, *args, **kwargs):
        acaras = AcaraModel.objects.all();

        return render(request, self.template_name, {'acaras': acaras})


class Bagian(LoginRequiredMixin, View):
    login_url = settings.LOGOUT_REDIRECT_URL
    template_name = 'bagian.html'

    def get(self, request, *args, **kwargs):
        bagians = BagianModel.objects.filter(acara_id=kwargs['id'])

        return render(request, self.template_name, {'bagians': bagians})


class Tugas(LoginRequiredMixin, View):
    login_url = settings.LOGOUT_REDIRECT_URL
    template_name = 'tugas.html'

    def get(self, request, *args, **kwargs):
        tasks = TugasModel.objects.filter(bagian_id=kwargs['id'])

        return render(request, self.template_name, {'tasks': tasks, 'bagian_id':kwargs['id']})


class EditTugas(LoginRequiredMixin, View):
    login_url = settings.LOGOUT_REDIRECT_URL
    template_name = 'edit-tugas.html'

    def post(self, request, *args, **kwargs):
        tugas = TugasModel.objects.filter(id=kwargs['id']).first()
        status_lama = tugas.status
        tugas.status = request.POST['status']
        tugas.komentar = request.POST['komentar']
        updated = 'updated'
        if len(request.FILES) != 0:
            if request.FILES['attachment_file'].size <= 2000000:
                tugas.attachment_file = request.FILES['attachment_file']
            else:
                updated = 'not-updated'
                return render(request, self.template_name, {'tugas': tugas, 'updated': updated})
        tugas.save()
        if status_lama == 'pending' and tugas.status == 'done':
            subyek = 'Tugas Diperbarui'
            pesan = request.user.username + ' telah memperbarui status dari ' + tugas.judul + '.'
            send_mail(subyek, pesan, settings.EMAIL_HOST_USER, [settings.ADMIN_EMAIL], fail_silently=False)

        return render(request, self.template_name, {'tugas': tugas, 'updated':updated})


    def get(self, request, *args, **kwargs):
        tugas = TugasModel.objects.filter(id=kwargs['id']).first()

        return render(request, self.template_name, {'tugas': tugas})

class PrintAcara(LoginRequiredMixin, View):
    login_url = settings.LOGOUT_REDIRECT_URL
    template_name = 'acara.html'

    def printToDocx(self, acara, bagians):
        doc = docx.Document();
        doc.add_heading(acara.judul, 0)
        doc.add_paragraph(acara.deskripsi)
        for bagian in bagians:
            doc.add_heading(bagian.divisi, level=1)
            doc.add_paragraph(bagian.deskripsi)
            tugas = TugasModel.objects.filter(bagian=bagian)
            tuga_table = doc.add_table(rows=1, cols=6)
            tuga_table.style = 'Table Grid'
            header_cells = tuga_table.rows[0].cells
            header_cells[0].text = 'No.'
            header_cells[1].text = 'Judul'
            header_cells[2].text = 'Tgl Dibuat'
            header_cells[3].text = 'Batas Tgl'
            header_cells[4].text = 'Tugas untuk'
            header_cells[5].text = 'Status'
            row_number = 0
            for tuga in tugas:
                row_cells = tuga_table.add_row().cells
                row_number += 1
                row_cells[0].text = str(row_number)
                row_cells[1].text = tuga.judul
                row_cells[2].text = str(tuga.tgl_dibuat)
                row_cells[3].text = str(tuga.batas_tgl)
                row_cells[4].text = str(tuga.tugas_untuk)
                row_cells[5].text = str(tuga.status)
        title = acara.judul + str(calendar.timegm(time.gmtime()))
        title = title + ".docx"
        doc.save("media/" + title)
        return title

    def get(self, request, *args, **kwargs):
        acara = AcaraModel.objects.filter(id=request.GET['acaraid']).first()
        bagians = BagianModel.objects.filter(acara=acara)
        file_name = self.printToDocx(acara, bagians)

        return HttpResponse(file_name)

